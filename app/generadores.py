from docx import Document
from openpyxl import Workbook
from docx import Document
from openpyxl import Workbook
from datetime import datetime
from config import EMPRESA, AUTOR

def crear_documentos(noticias, sufijo):
    fecha_hora = datetime.now().strftime("%Y-%m-%d_%H-%M")

    doc = Document()
    wb = Workbook()
    ws = wb.active

    doc.add_heading(EMPRESA, level=1)
    doc.add_paragraph(f"Fecha: {fecha_hora}")
    doc.add_paragraph(f"Autor: {AUTOR}")

    doc.add_heading("Resumen Ejecutivo", level=2)
    doc.add_paragraph(
        "Documento generado automáticamente como parte del proceso de análisis periódico."
    )

    ws.append(["Título", "Resumen"])

    doc.add_heading("Noticias Relevantes", level=2)

    for i, noticia in enumerate(noticias, 1):
        doc.add_heading(f"{i}. {noticia['titulo']}", level=3)
        doc.add_paragraph(noticia["resumen"])
        ws.append([noticia["titulo"], noticia["resumen"]])

    ruta_word = f"output/reportes/reporte_{sufijo}_{fecha_hora}.docx"
    ruta_excel = f"output/reportes/reporte_{sufijo}_{fecha_hora}.xlsx"

    doc.save(ruta_word)
    wb.save(ruta_excel)

    return ruta_word, ruta_excel